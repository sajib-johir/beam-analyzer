import io
import math
import numpy as np
import streamlit as st
import matplotlib
matplotlib.use("Agg")  # headless
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import pandas as pd

st.set_page_config(page_title="Mechanical Beam Analyzer", layout="wide", initial_sidebar_state="expanded")

def unit_tokens(u: str):
    if u.startswith("Imperial"):
        return {"len":"ft","force":"lbf","moment":"lbf·ft","I":"in^4","len_small":"ft"}
    if u.startswith("CGS"):
        return {"len":"cm","force":"dyn","moment":"dyn·cm","I":"cm^4","len_small":"cm"}
    return {"len":"m","force":"N","moment":"Nm","I":"m^4","len_small":"m"}

# ------------------------- Built-in materials (expanded) -------------------------
BUILTIN_MATERIALS = {
    # --- Structural Steels (ASTM / EN / BS) ---
    "ASTM A36": (2.00e11, 2.50e8),
    "ASTM A572 Grade 50": (2.00e11, 3.45e8),
    "ASTM A588": (2.00e11, 3.45e8),
    "ASTM A992": (2.00e11, 3.45e8),
    "ASTM A913 Grade 50": (2.00e11, 3.45e8),
    "ASTM A913 Grade 65": (2.00e11, 4.50e8),
    "ASTM A514 (T-1)": (2.00e11, 6.90e8),
    "ASTM A709 Grade 50": (2.00e11, 3.45e8),
    "EN 10025 S235": (2.10e11, 2.35e8),
    "EN 10025 S275": (2.10e11, 2.75e8),
    "EN 10025 S355": (2.10e11, 3.55e8),
    "EN 10025 S460": (2.10e11, 4.60e8),
    "BS 4360 43A": (2.00e11, 2.75e8),
    "BS 4360 50D": (2.00e11, 3.55e8),

    # --- Alloy Steels ---
    "AISI 1045": (2.05e11, 5.30e8),
    "AISI 4140 (QT)": (2.10e11, 6.55e8),
    "AISI 4340 (QT)": (2.05e11, 8.50e8),
    "AISI 8620": (2.05e11, 3.80e8),

    # --- Stainless / Duplex ---
    "Stainless 304": (1.93e11, 2.15e8),
    "Stainless 304L": (1.93e11, 1.70e8),
    "Stainless 316": (1.93e11, 2.90e8),
    "Stainless 316L": (1.93e11, 1.70e8),
    "Stainless 410": (2.00e11, 2.75e8),
    "Stainless 420": (2.00e11, 4.50e8),
    "Stainless 904L": (1.93e11, 2.20e8),
    "Stainless 17-4PH": (1.96e11, 1.10e9),
    "Duplex 2205": (2.00e11, 4.50e8),
    "Super Duplex 2507": (2.00e11, 5.50e8),

    # --- Aluminum Alloys ---
    "Aluminum 2024-T3": (7.30e10, 3.24e8),
    "Aluminum 5052-H32": (7.00e10, 1.93e8),
    "Aluminum 5083-H116": (7.00e10, 2.15e8),
    "Aluminum 6061-T6": (6.90e10, 2.70e8),
    "Aluminum 6063-T5": (6.90e10, 1.45e8),
    "Aluminum 6063-T6": (6.90e10, 2.20e8),
    "Aluminum 6082-T6": (6.90e10, 2.50e8),
    "Aluminum 7075-T6": (7.20e10, 5.00e8),
    "Aluminum 7075-T73": (7.10e10, 4.35e8),

    # --- Titanium / Nickel / Cobalt / Magnesium ---
    "Titanium Grade 2": (1.05e11, 2.75e8),
    "Titanium Grade 9": (1.05e11, 6.20e8),
    "Titanium Ti-6Al-4V (Gr5)": (1.14e11, 8.80e8),
    "Inconel 625": (2.05e11, 4.60e8),
    "Inconel 718": (2.10e11, 1.00e9),
    "Hastelloy C-276": (2.05e11, 2.83e8),
    "Monel 400": (1.80e11, 1.70e8),
    "Cobalt-Chrome": (2.30e11, 6.00e8),
    "Magnesium AZ31B": (4.50e10, 2.00e8),
    "Magnesium ZK60A": (4.50e10, 2.60e8),

    # --- Copper / Brass / Bronze ---
    "Copper C110": (1.10e11, 7.00e7),
    "Cu-Ni 70/30": (1.50e11, 1.40e8),
    "Brass C260 (70/30)": (9.70e10, 9.50e7),
    "Brass 360": (1.00e11, 2.10e8),
    "Bronze C93200": (1.00e11, 1.70e8),
    "Phosphor Bronze C51000": (1.10e11, 2.75e8),

    # --- Polymers ---
    "ABS": (2.00e9, 3.50e7),
    "PLA": (3.50e9, 6.00e7),
    "PETG": (2.20e9, 5.00e7),
    "Nylon 6/6": (2.80e9, 8.00e7),
    "POM (Delrin)": (3.00e9, 7.00e7),
    "Polycarbonate": (2.30e9, 7.00e7),
    "PEEK": (3.60e9, 9.00e7),
    "PTFE (Teflon)": (5.00e8, 2.00e7),
    "UHMWPE": (8.00e8, 2.00e7),
    "PVC (rigid)": (3.00e9, 5.50e7),

    # --- Woods & Engineered Wood ---
    "Douglas Fir-Larch": (1.20e10, 5.00e7),
    "Oak": (1.20e10, 4.00e7),
    "Pine": (9.00e9, 3.50e7),
    "Maple": (1.20e10, 4.50e7),
    "Birch": (1.10e10, 4.50e7),
    "Spruce-Pine-Fir": (9.50e9, 3.50e7),
    "Glulam (24F-V4)": (1.20e10, None),
    "LVL (2.0E)": (1.38e10, None),
    "Southern Pine No.2": (1.10e10, None),

    # --- Composites ---
    "CFRP (Uni, 60% Vf)": (1.35e11, None),
    "GFRP (Uni, 55% Vf)": (2.50e10, None),

    # --- Reinforced Concrete (effective, depends on rebar) ---
    "Reinforced Concrete": (3.00e10, None),
}

# ------------------------- Helpers -------------------------
def sci_notation(val: float) -> str:
    try:
        f = float(val)
        if f == 0:
            return "0"
        exp = int(math.floor(math.log10(abs(f))))
        base = f / (10**exp)
        return f"{base:.3f}*10^{exp}"
    except Exception:
        return str(val)

def validate_inputs(L, a, I, c, E):
    problems = []
    if L is None or L <= 0: problems.append("Beam length L must be > 0.")
    if I is None or I <= 0: problems.append("Area moment I must be > 0.")
    if c is None or c <= 0: problems.append("Extreme fiber distance c must be > 0.")
    if E is None or E <= 0: problems.append("Young's Modulus E must be > 0.")
    if a is not None and (a < 0): problems.append("Load position a must be ≥ 0.")
    if a is not None and L is not None and a > L: problems.append("Load position a must be ≤ L.")
    return problems

def pad_limits(values, frac=0.12):
    v = np.asarray(values, dtype=float)
    vmin = float(np.nanmin(v)) if v.size else 0.0
    vmax = float(np.nanmax(v)) if v.size else 0.0
    if vmin == vmax == 0.0:
        return (-1.0, 1.0)
    span = vmax - vmin
    if span == 0:
        pad = max(abs(vmax), 1.0) * frac
        return (vmin - pad, vmax + pad)
    pad = span * frac
    if vmin < 0 < vmax:
        m = max(abs(vmin), abs(vmax))
        return (-m - pad, m + pad)
    return (vmin - pad, vmax + pad)

# ------------------------- Beam model -------------------------
class Beam:
    def __init__(self, beam_type, load_type, length, load, E, I, c, material="Custom", yield_strength=None, a=None):
        self.beam_type = beam_type
        self.load_type = load_type
        self.length = float(length)
        self.load = float(load)
        self.E = float(E)
        self.I = float(I)
        self.c = float(c)
        self.material = material
        self.yield_strength = (None if yield_strength in (None, "", "None") else float(yield_strength))
        self.a = (None if a in (None, "", "None") else float(a))

    def diagrams(self, n=601):
        L = self.length
        x = np.linspace(0.0, L, n)
        P_or_w = self.load
        b = self.beam_type
        case = self.load_type
        V = np.zeros_like(x)
        M = np.zeros_like(x)

        if b == "Simply Supported":
            if case == "Point Load (Center)":
                P = P_or_w; R = P/2
                V = np.where(x < L/2, R, R - P)
                M = np.where(x <= L/2, R*x, R*x - P*(x - L/2))
            elif case == "Point Load (Any Position)":
                P = P_or_w; a = self.a if self.a is not None else L/2
                R1 = P * (L - a) / L
                V = np.where(x < a, R1, R1 - P)
                M = np.where(x < a, R1*x, R1*x - P*(x - a))
            elif case == "UDL (Uniformly Distributed Load)":
                w = P_or_w; R = w*L/2
                V = R - w*x
                M = R*x - 0.5*w*x**2
            elif case == "Applied Moment (Center)":
                M0 = P_or_w
                V = np.zeros_like(x)
                M = np.where(x < L/2, 0.0, M0)

        elif b == "Cantilever":
            if case == "Point Load (End)":
                P = P_or_w
                V = -P * np.ones_like(x)
                M = -P * (L - x)
            elif case == "UDL (Uniformly Distributed Load)":
                w = P_or_w
                V = -w*(L - x)
                M = -0.5*w*(L - x)**2
            elif case == "Applied Moment (End)":
                M0 = P_or_w
                V = np.zeros_like(x)
                M = np.full_like(x, M0)

        elif b == "Fixed-Fixed":
            if case == "UDL (Uniformly Distributed Load)":
                w = P_or_w
                R = w*L/2
                M_A = -w*L**2/12.0
                V = R - w*x
                M = R*x - 0.5*w*x**2 + M_A
            elif case == "Point Load (Center)":
                P = P_or_w
                R = P/2
                M_A = -P*L/8.0
                M = np.where(x <= L/2, R*x + M_A, R*x - P*(x - L/2) + M_A)
                V = np.where(x < L/2, R, R - P)

        return x, V, M

    def deflection(self, x, M):
        EI = self.E * self.I
        if EI == 0: return np.zeros_like(x)
        k = M / EI
        I1 = np.cumsum(np.concatenate([[0.0], 0.5*(k[1:] + k[:-1]) * np.diff(x)]))
        I2 = np.cumsum(np.concatenate([[0.0], 0.5*(I1[1:] + I1[:-1]) * np.diff(x)]))
        # boundary conditions similar to desktop
        if self.beam_type == "Cantilever":
            C1, C2 = 0.0, 0.0
        else:
            C2 = 0.0
            C1 = -I2[-1] / (x[-1] if x[-1] != 0 else 1.0)
        return I2 + C1*x + C2

    def reactions(self):
        L = self.length; case = self.load_type; b = self.beam_type; d = {}
        if b == "Simply Supported":
            if case == "Point Load (Center)":
                P = self.load; d["RA"] = P/2; d["RB"] = P/2
            elif case == "Point Load (Any Position)":
                P = self.load; a = self.a if self.a is not None else L/2
                d["RA"] = P*(L - a)/L; d["RB"] = P*a/L
            elif case == "UDL (Uniformly Distributed Load)":
                w = self.load; W = w*L; d["RA"] = W/2; d["RB"] = W/2
            elif case == "Applied Moment (Center)":
                M0 = self.load; d["MA@A"] = M0/2; d["MB@B"] = -M0/2
        elif b == "Cantilever":
            if case == "Point Load (End)":
                P = self.load; d["V@fix"] = P; d["M@fix"] = P*L
            elif case == "UDL (Uniformly Distributed Load)":
                w = self.load; d["V@fix"] = w*L; d["M@fix"] = w*L**2/2
            elif case == "Applied Moment (End)":
                M0 = self.load; d["V@fix"] = 0.0; d["M@fix"] = M0
        elif b == "Fixed-Fixed":
            if case == "UDL (Uniformly Distributed Load)":
                w = self.load
                d["RA"] = w*L/2; d["RB"] = w*L/2
                d["MA@A"] = w*L**2/12; d["MB@B"] = -w*L**2/12
            elif case == "Point Load (Center)":
                P = self.load
                d["RA"] = P/2; d["RB"] = P/2
                d["MA@A"] = P*L/8; d["MB@B"] = -P*L/8
        return d

    def bending_stress(self):
        x, _, M = self.diagrams()
        Mabs = np.max(np.abs(M))
        if self.I == 0 or self.c is None:
            return None
        return (Mabs * self.c) / self.I

    def safety_factor(self):
        s = self.bending_stress()
        if self.yield_strength and s not in (None, 0):
            return self.yield_strength / s
        return None

# ------------------------- Theme & Styling -------------------------
# Theme toggle (light/dark)
# Locked to dark theme


css_extra = """
/* Title spacing */
h1 { margin-top: 0.6rem !important; margin-bottom: 0.6rem !important; }
/* App & structure */
.block-container { padding-top: 0.6rem !important; }
div[data-testid="stSidebar"] { background: var(--card) !important; border-right: 1px solid var(--border) !important; }
div[data-testid="stSidebar"] * { color: var(--text) !important; }
/* Headings & text */
h1,h2,h3,h4,h5,h6, p, span, label, small, strong, em, .stMarkdown, .stCaption, .stText { color: var(--text) !important; }
/* Cards / expanders / alerts / tabs */
.st-expander, .stTabs, .stDataFrame, .stAlert, .stToast, div[role="tablist"], div[role="tab"], div[role="alert"] {
    background: var(--card) !important; color: var(--text) !important; border: 1px solid var(--border) !important; border-radius: 12px !important;
}
.st-expander summary { color: var(--text) !important; }
/* Inputs */
.stTextInput input, .stNumberInput input, .stTextArea textarea {
    background: var(--input-bg) !important; color: var(--input-text) !important;
    border: 1px solid var(--input-border) !important; border-radius: 10px !important;
}


input::placeholder, textarea::placeholder { color: var(--placeholder) !important; opacity: 1; }
/* Dropdowns (Selectbox) */
.stSelectbox div[data-baseweb="select"] {
    background: var(--input-bg) !important;
    color: var(--input-text) !important;
    border: 1px solid var(--input-border) !important;
    border-radius: 10px !important;
}
.stSelectbox div[data-baseweb="select"]:hover {
    border: 1px solid #3ba891 !important;
}
.stSelectbox * { color: var(--input-text) !important; }

/* File uploader (entire content) */
div[data-testid="stFileUploader"] { border:2px solid var(--border) !important;
    border-style: solid !important; border-width: 2px !important;
    background: var(--card) !important; color: var(--text) !important; 
    border: 1px dashed var(--border) !important; border-radius: 12px !important;
}
div[data-testid="stFileUploader"] * { color: var(--text) !important; }
/* Buttons */
.stButton>button {
    background-color: #4cbfa6 !important;
    border: 2px solid #3ba891 !important;
    color: #ffffff !important;
    font-weight: 600 !important;
    border-radius: 14px !important;
    padding: 0.6rem 1rem !important;
    box-shadow: 0 3px 14px rgba(0,0,0,.18) !important;
    cursor: pointer !important;
    transition: background-color .15s ease, border-color .15s ease, transform .08s ease;
}
.stButton>button:hover {
    background-color: #3ba891 !important;
    border: 2px solid #3ba891 !important;
    transform: translateY(-1px);
}
.stButton>button:active { transform: translateY(0) !important; box-shadow: 0 1px 6px rgba(0,0,0,.20) !important; }
/* Tables / code */
pre, code, .stCode { background: var(--card) !important; color: var(--text) !important; border-radius: 10px !important; }
/* Radio/checkbox bullets */
div[role="radiogroup"] *, div[role="checkbox"] * { color: var(--text) !important; }
/* Section headers */
hr, .stDivider { border-color: var(--border) !important; }
"""

st.markdown("<style>" + css_extra + "</style>", unsafe_allow_html=True)


# ------------------------- Title -------------------------
st.title("Mechanical Beam Analyzer")
st.caption("E = Young's Modulus (Pa) • Sy = Yield Strength (Pa). Upload CSV with headers: name,E_Pa,Sy_Pa")

# ------------------------- Sidebar (Material first) -------------------------
with st.sidebar:
    with st.form("inputs_form", clear_on_submit=False):
        st.header("Material")
        st.write("**CSV headers must be:** `name,E_Pa,Sy_Pa`.")
        uploaded = st.file_uploader("Import materials CSV", type=["csv"])  # users can add more materials
        if st.session_state.get('loaded_json'):
            st.info('JSON loaded — click to apply values below.')

        # Merge built-ins + CSV (dedupe by name)
        materials = dict(BUILTIN_MATERIALS)
        imported_preview = None
        if uploaded is not None:
            try:
                df_csv = pd.read_csv(uploaded)
                added = 0
                for _, row in df_csv.iterrows():
                    name = str(row[0]).strip()
                    E = None if pd.isna(row[1]) else float(row[1])
                    Sy = None if pd.isna(row[2]) else float(row[2])
                    if name:
                        materials[name] = (E, Sy)
                        added += 1
                imported_preview = df_csv.head()
                st.success(f"Added {added} materials from CSV.")
            except Exception as e:
                st.error(f"Could not read CSV: {e}")

        # Categorize + Code bundle tags
        def code_bundle_of(name:str):
            n = name.lower()
            tags = []
            if any(k in n for k in ["astm","a36","a572","a588","a992","a913","a514","a709"]): tags.append("AISC/ASTM")
            if any(k in n for k in ["en 10025","s235","s275","s355","s355j2"]): tags.append("Eurocode EN")
            if "bs 4360" in n: tags.append("BS (UK)")
            if any(k in n for k in ["6061","6063","6082","2024","5052","5083","7075","aluminum"]): tags.append("Aluminum Assoc.")
            if any(k in n for k in ["glulam","lvl","southern pine","spf","fir","larch"]): tags.append("NDS Timber")
            if any(k in n for k in ["cfrp","gfrp","frp","carbon","glass"]): tags.append("Composites")
            return tags or ["General"]
        def category_of(name:str):
            n = name.lower()
            if any(k in n for k in ["a36","a572","a588","a992","a913","a514","a709","s235","s275","s355","bs 4360","steel"]): return "Steel"
            if any(k in n for k in ["aluminum","6061","6063","6082","2024","5052","5083","7075"]): return "Aluminum"
            if any(k in n for k in ["wood","glulam","lvl","southern pine","spf","fir","larch"]): return "Timber"
            if any(k in n for k in ["cfrp","gfrp","frp","carbon","glass"]): return "Composite"
            if any(k in n for k in ["titanium","magnesium","monel","inconel","bronze","copper","phosphor"]): return "Other Metal"
            if any(k in n for k in ["stainless","duplex"]): return "Stainless"
            return "Other"

        # Show mode: All vs Preset Packs
        show_mode = st.radio("Material view", ["All materials", "Preset packs"], horizontal=True)

        # If preset packs, let user pick packs (can combine)
        bundle_options = ["AISC/ASTM","Eurocode EN","BS (UK)","Aluminum Assoc.","NDS Timber","Composites"]
        selected_packs = []
        if show_mode == "Preset packs":
            selected_packs = st.multiselect("Choose packs", options=bundle_options, default=["AISC/ASTM"]) 

        # Search box (works for both modes)
        mat_search = st.text_input("Search material… (type to filter)")

        # Build filtered list based on mode + search
        all_names = sorted(materials.keys())
        filt = []
        for m in all_names:
            if mat_search and mat_search.lower() not in m.lower():
                continue
            if show_mode == "Preset packs":
                packs = code_bundle_of(m)
                if not selected_packs or not any(p in packs for p in selected_packs):
                    continue
            filt.append(m)

        if not filt:
            st.warning("No materials match the current filters.")
            filt = ["(none)"]

        mat_name = st.selectbox("Select material", options=filt)
        E_default, Sy_default = materials.get(mat_name, (2.10e11, None))

        # Badges for selected material
        cat = category_of(mat_name)
        badges = "".join([f"<span class='badge'>{b}</span>" for b in code_bundle_of(mat_name)])
        st.markdown(f"**Selected:** {mat_name} <span class='badge muted'>{cat}</span> {badges}", unsafe_allow_html=True)

        E = st.number_input("E — Young's Modulus [Pa]", value=float(E_default or 2.10e11), step=1e9, format="%.6e")
        Sy = st.text_input("Sy — Yield Strength [Pa] (optional)", value=("" if Sy_default is None else f"{Sy_default:.6e}"))

        st.divider()
        st.header("Beam Setup")
        units = st.radio("Units", ["Metric (SI)", "Imperial", "CGS"], horizontal=True, index=0)
        beam_type = st.selectbox("Beam Type", ["Simply Supported", "Cantilever", "Fixed-Fixed"])
        load_options = {
            "Simply Supported": ["Point Load (Center)","Point Load (Any Position)","UDL (Uniformly Distributed Load)","Applied Moment (Center)"],
            "Cantilever": ["Point Load (End)","UDL (Uniformly Distributed Load)","Applied Moment (End)"],
            "Fixed-Fixed": ["UDL (Uniformly Distributed Load)", "Point Load (Center)"],
        }
        load_type = st.selectbox("Load Type", load_options[beam_type])

        L_label = "L — Beam Length [m]" if units.startswith("Metric") else ("L — Beam Length [ft]" if units.startswith("Imperial") else "L — Beam Length [cm]")
        L = st.number_input(L_label, min_value=0.01, value=5.0, step=0.1, format="%.3f")

        if "UDL" in load_type:
            load_label = "Load value [N/m]" if units.startswith("Metric") else ("Load value [lbf/ft]" if units.startswith("Imperial") else "Load value [dyn/cm]")
        elif "Moment" in load_type:
            load_label = "Load value [Nm]" if units.startswith("Metric") else ("Load value [lbf·ft]" if units.startswith("Imperial") else "Load value [dyn·cm]")
        else:
            load_label = "Load value [N]" if units.startswith("Metric") else ("Load value [lbf]" if units.startswith("Imperial") else "Load value [dyn]")
        load_val = st.number_input(load_label, value=12000.0, step=100.0, format="%.3f")

        a = None
        if "Any Position" in load_type:
            a_label = "a — Point load position [m]" if units.startswith("Metric") else ("a — Point load position [ft]" if units.startswith("Imperial") else "a — Point load position [cm]")
            a = st.number_input(a_label, min_value=0.0, value=float(L/3), step=0.05, format="%.3f")

        preset = st.selectbox("Quick preset", ["(none)", "SS + Center Point", "Cantilever + End Point", "SS + UDL"]) 
        applied = st.form_submit_button("Apply inputs (Enter)")
        if applied and preset != "(none)":
            if preset == "SS + Center Point":
                beam_type = "Simply Supported"; load_type = "Point Load (Center)"; L = 4.0; load_val = 8000.0
            elif preset == "Cantilever + End Point":
                beam_type = "Cantilever"; load_type = "Point Load (End)"; L = 2.0; load_val = 1200.0
            elif preset == "SS + UDL":
                beam_type = "Simply Supported"; load_type = "UDL (Uniformly Distributed Load)"; L = 5.0; load_val = 3000.0
            st.info("Preset applied. You can now run Results/Plots.")

# ------------------------- Section Library (auto I & c) ------------------------- (auto I & c) ------------------------- (auto I & c) ------------------------- (auto I & c) -------------------------
SECTION_SPECS = {
    "Rectangle (b,h)": (["b (width) [m]", "h (height) [m]"], "I = b·h^3/12,  c = h/2"),
    "Square (b)": (["b (side) [m]"], "I = b^4/12,  c = b/2"),
    "Solid Circle (D)": (["D (diameter) [m]"], "I = π·D^4/64,  c = D/2"),
    "Hollow Circle (Do,Di)": (["Do (outer Ø) [m]", "Di (inner Ø) [m]"], "I = (π/64)(Do^4 − Di^4),  c = Do/2"),
    "Rectangular Tube (b,h,t)": (["b (outer) [m]", "h (outer) [m]", "t (wall) [m]"], "I = [b·h^3 − (b−2t)(h−2t)^3]/12,  c = h/2"),
    "I-Beam (bf,tf,tw,h)": (["bf [m]","tf [m]","tw [m]","h [m]"], "c = h/2; hw = h−2tf; I = 2[bf·tf^3/12 + bf·tf(h/2 − tf/2)^2] + (tw·hw^3)/12"),
}

def compute_Ic(shape, dims):
    if shape == "Rectangle (b,h)":
        b, h = dims
        I = b*h**3/12.0; c = h/2.0
    elif shape == "Square (b)":
        (b,) = dims
        I = b**4/12.0; c = b/2.0
    elif shape == "Solid Circle (D)":
        (D,) = dims
        I = (math.pi*D**4)/64.0; c = D/2.0
    elif shape == "Hollow Circle (Do,Di)":
        Do, Di = dims
        if Di >= Do: raise ValueError("Inner diameter must be smaller than outer diameter.")
        I = (math.pi*(Do**4 - Di**4))/64.0; c = Do/2.0
    elif shape == "Rectangular Tube (b,h,t)":
        b, h, t = dims
        if t <= 0 or b <= 2*t or h <= 2*t: raise ValueError("Need b>2t and h>2t.")
        I = (b*h**3 - (b-2*t)*(h-2*t)**3)/12.0; c = h/2.0
    elif shape == "I-Beam (bf,tf,tw,h)":
        bf, tf, tw, h = dims
        hw = h - 2*tf
        if hw <= 0 or tw <= 0: raise ValueError("Require h > 2tf and tw > 0.")
        Ifl = (bf*tf**3)/12.0 + bf*tf*((h/2 - tf/2)**2)
        Ifr = (bf*tf**3)/12.0 + bf*tf*((h/2 - tf/2)**2)
        Iw  = (tw*hw**3)/12.0
        I = Ifl + Ifr + Iw; c = h/2.0
    else:
        raise ValueError("Unsupported shape")
    return float(I), float(c)

st.subheader("Section Library (auto I & c) — or enter manually below")
col1, col2 = st.columns([1,2])
with col1:
    shape = st.selectbox("Shape", ["(none)"] + list(SECTION_SPECS.keys()))
    dims_vals = []
    if shape != "(none)":
        labels, help_txt = SECTION_SPECS[shape]
        for lab in labels:
            dims_vals.append(st.number_input(lab, min_value=0.0, value=0.05 if "b" in lab or "Do" in lab or "bf" in lab else 0.03, step=0.001, format="%.6f", key=lab))
        st.caption(help_txt)
        if st.button("Compute I & c from shape"):
            try:
                I_auto, c_auto = compute_Ic(shape, tuple(dims_vals))
                st.session_state["I_auto"] = I_auto
                st.session_state["c_auto"] = c_auto
                st.success(f"Computed  I = {I_auto:.6e} m⁴,  c = {c_auto:.6e} m")
            except Exception as ex:
                st.error(str(ex))

with col2:
    I = st.number_input(f"I — Second Moment of Area [{unit_tokens(units)['I']}]", min_value=0.0, value=float(st.session_state.get("I_auto", 1.125e-7)), step=1e-10, format="%.6e")
    c = st.number_input(f"c — Distance to Extreme Fiber [{unit_tokens(units)['len_small']}]", min_value=0.0, value=float(st.session_state.get("c_auto", 0.015)), step=1e-4, format="%.6e")
    st.caption("σ = M·c / I • Higher I → less deflection")

# ------------------------- Separate action buttons -------------------------
st.subheader("Run / Export")
btn1, btn2, btn3, btn4, btn5, btn6 = st.columns([1,1,1,1,1,1])
with btn1:
    do_results = st.button("Show Results")
with btn2:
    do_plots = st.button("Show Plots")
with btn3:
    do_pdf = st.button("Download PDF")
with btn4:
    do_csv = st.button("Download CSV")
with btn5:
    do_xlsx = st.button("Export Excel")
with btn6:
    do_docx = st.button("Export Word")

# JSON save/load row
btn7, btn8 = st.columns([1,1])
with btn7:
    do_save_json = st.button("Save Inputs (JSON)")
with btn8:
    up_json = st.file_uploader("Load Inputs (JSON)", type=["json"], key="load_json_upl", help="Load a previously saved analysis (material, units, geometry, loads, section). Nothing is applied until you confirm.")


# Toggle interactive plots
interactive = st.toggle("Interactive plots (Plotly)", value=False)

# ------------------------- Compute on demand (cached) -------------------------
@st.cache_data(show_spinner=False)
def compute_all(beam_type, load_type, L, load_val, E, I, c, material, Sy, a):
    beam = Beam(beam_type, load_type, L, load_val, E, I, c, material=material, yield_strength=Sy, a=a)
    x, V, M = beam.diagrams()
    y = beam.deflection(x, M)
    stress = beam.bending_stress()
    sf = beam.safety_factor()
    react = beam.reactions()
    return beam, x, V, M, y, stress, sf, react

# ------------------------- Panels -------------------------
if do_results:
    try:
        errs = validate_inputs(L, a, I, c, E)
        if errs:
            st.error("\n".join(errs))
            raise Exception("ValidationError")
        # unit conversions to SI for computation
        L_si = L * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        if "UDL" in load_type:
            load_si = load_val * (14.5939 if units.startswith("Imperial") else (1e-3 if units.startswith("CGS") else 1.0))  # lbf/ft→N/m; dyn/cm→N/m
        elif "Moment" in load_type:
            load_si = load_val * (1.355818 if units.startswith("Imperial") else (1e-7 if units.startswith("CGS") else 1.0))  # lbf·ft→Nm; dyn·cm→Nm
        else:
            load_si = load_val * (4.4482216153 if units.startswith("Imperial") else (1e-5 if units.startswith("CGS") else 1.0))  # lbf→N; dyn→N
        a_si = None if a is None else a * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        E_si = E  # keep Pa
        I_si = I
        c_si = c

        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L_si, load_si, E_si, I_si, c_si, mat_name, Sy, a_si)

        # Display in chosen units
        def to_units_V(v):
            return (v / 4.4482216153) if units.startswith("Imperial") else ((v * 1e5) if units.startswith("CGS") else v)
        def to_units_M(v):
            return (v / 1.355818) if units.startswith("Imperial") else ((v * 1e7) if units.startswith("CGS") else v)
        def to_units_x(v):
            return (v / 0.3048) if units.startswith("Imperial") else ((v * 100.0) if units.startswith("CGS") else v)
        def to_units_defl(v):
            return (v / 0.3048) if units.startswith("Imperial") else ((v * 100.0) if units.startswith("CGS") else v)

        st.subheader("Results")
        res = []
        res.append(f"Beam Type: {beam.beam_type}")
        res.append(f"Load Type: {beam.load_type}")
        if 'Any Position' in beam.load_type and beam.a is not None:
            res.append(f"Point Load Position a: {to_units_x(beam.a):.4f} {'ft' if units.startswith('Imperial') else ('cm' if units.startswith('CGS') else 'm')}")
        res.append(f"Length L: {to_units_x(beam.length):.4f} {'ft' if units.startswith('Imperial') else ('cm' if units.startswith('CGS') else 'm')}")
        res.append(f"Load Value: {sci_notation(load_val)} {'(lbf, lbf/ft, or lbf·ft)' if units.startswith('Imperial') else ('(dyn, dyn/cm, or dyn·cm)' if units.startswith('CGS') else '(N, N/m, or Nm)')}")
        res.append(f"Material: {beam.material}")
        res.append(f"Young's Modulus E: {(sci_notation(E*10)+' dyn/cm²') if units.startswith('CGS') else (sci_notation(E)+' Pa')}")
        if Sy: res.append(f"Yield Strength Sy: {(sci_notation(Sy*10)+' dyn/cm²') if units.startswith('CGS') else (sci_notation(Sy)+' Pa')}")
        res.append(f"I (area moment): {sci_notation(I)} m^4")
        res.append(f"c (outer fiber dist.): {sci_notation(c)} m")
        res.append("")
        res.append(f"Max |V| (Shear): {sci_notation(np.max(np.abs(to_units_V(V))))} {'lbf' if units.startswith('Imperial') else ('dyn' if units.startswith('CGS') else 'N')}")
        res.append(f"Max |M| (Bending): {sci_notation(np.max(np.abs(to_units_M(M))))} {'lbf·ft' if units.startswith('Imperial') else ('dyn·cm' if units.startswith('CGS') else 'Nm')}")
        res.append(f"Max |δ| (Deflection): {sci_notation(np.max(np.abs(to_units_defl(y))))} {'ft' if units.startswith('Imperial') else ('cm' if units.startswith('CGS') else 'm')}")
        res.append(f"Max σ (from max |M|): {((sci_notation(stress*10)+' dyn/cm²') if (stress is not None and units.startswith('CGS')) else ((sci_notation(stress)+' Pa') if stress is not None else 'N/A'))}")
        if sf: res.append(f"Safety Factor n = Sy/σ: {sf:.2f}")
        if react:
            res.append("")
            res.append("Support Reactions (SI units):")
            for k, v in react.items():
                res.append(f"  {k}: {sci_notation(v)}")
        if sf and sf < 1.5:
            res.append("")
            res.append("Note: Low safety factor (<1.5). Consider increasing I, reducing load, shortening span, or choosing a stronger material.")
        st.code("\n".join(res))
    except Exception as e:
        st.error(str(e))

if do_plots:
    try:
        errs = validate_inputs(L, a, I, c, E)
        if errs:
            st.error("\n".join(errs))
            raise Exception("ValidationError")
        # recompute in SI
        L_si = L * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        if "UDL" in load_type:
            load_si = load_val * (14.5939 if units.startswith("Imperial") else (1e-3 if units.startswith("CGS") else 1.0))
        elif "Moment" in load_type:
            load_si = load_val * (1.355818 if units.startswith("Imperial") else (1e-7 if units.startswith("CGS") else 1.0))
        else:
            load_si = load_val * (4.4482216153 if units.startswith("Imperial") else (1e-5 if units.startswith("CGS") else 1.0))
        a_si = None if a is None else a * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        E_si = E; I_si = I; c_si = c
        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L_si, load_si, E_si, I_si, c_si, mat_name, Sy, a_si)

        if interactive:
            import plotly.graph_objects as go
            tab1, tab2, tab3 = st.tabs(["Shear", "Moment", "Deflection"])
            with tab1:
                fig = go.Figure(); fig.add_trace(go.Scatter(x=x, y=V, mode='lines', name='V(x) [N]'))
                imax_v = int(np.argmax(np.abs(V))); fig.add_trace(go.Scatter(x=[x[imax_v]], y=[V[imax_v]], mode='markers+text', text=["max |V|"], textposition="top center", name="max |V|"))
                fig.update_layout(title='Shear V', xaxis_title='x (m)', yaxis_title='V (N)'); st.plotly_chart(fig, use_container_width=True)
            with tab2:
                fig = go.Figure(); fig.add_trace(go.Scatter(x=x, y=M, mode='lines', name='M(x) [Nm]'))
                imax_m = int(np.argmax(np.abs(M))); fig.add_trace(go.Scatter(x=[x[imax_m]], y=[M[imax_m]], mode='markers+text', text=["max |M|"], textposition="top center", name="max |M|"))
                fig.update_layout(title='Moment M', xaxis_title='x (m)', yaxis_title='M (Nm)'); st.plotly_chart(fig, use_container_width=True)
            with tab3:
                fig = go.Figure(); fig.add_trace(go.Scatter(x=x, y=y, mode='lines', name='δ(x) [m]'))
                imax_y = int(np.argmax(np.abs(y))); fig.add_trace(go.Scatter(x=[x[imax_y]], y=[y[imax_y]], mode='markers+text', text=["max |δ|"], textposition="top center", name="max |δ|"))
                fig.update_layout(title='Deflection δ', xaxis_title='x (m)', yaxis_title='δ (m)'); st.plotly_chart(fig, use_container_width=True)
        else:
            fig = plt.figure(figsize=(10, 10))
            fig.suptitle(f"{beam.beam_type} – {beam.load_type}", fontsize=14)
            ax1 = fig.add_subplot(3,1,1)
            ax1.plot(x, V, linewidth=2.2, label="Shear V(x) [N]")
            imax_v = int(np.argmax(np.abs(V))); ax1.plot([x[imax_v]], [V[imax_v]], marker="o"); ax1.annotate("max |V|", (x[imax_v], V[imax_v]))
            ax1.set_title("Shear Force"); ax1.set_ylabel("V (N)"); ax1.grid(True, alpha=0.3); ax1.legend(); ax1.axhline(0, linewidth=0.8)
            ax1.set_ylim(*pad_limits(V))
            ax2 = fig.add_subplot(3,1,2)
            ax2.plot(x, M, linewidth=2.2, label="Moment M(x) [Nm]")
            imax_m = int(np.argmax(np.abs(M))); ax2.plot([x[imax_m]], [M[imax_m]], marker="o"); ax2.annotate("max |M|", (x[imax_m], M[imax_m]))
            ax2.set_title("Bending Moment"); ax2.set_ylabel("M (Nm)"); ax2.grid(True, alpha=0.3); ax2.legend(); ax2.axhline(0, linewidth=0.8)
            ax2.set_ylim(*pad_limits(M))
            ax3 = fig.add_subplot(3,1,3)
            ax3.plot(x, y, linewidth=2.2, label="Deflection δ(x) [m]")
            imax_y = int(np.argmax(np.abs(y))); ax3.plot([x[imax_y]], [y[imax_y]], marker="o"); ax3.annotate("max |δ|", (x[imax_y], y[imax_y]))
            ax3.set_title("Deflection"); ax3.set_xlabel("x (m)"); ax3.set_ylabel("δ (m)")
            ax3.grid(True, alpha=0.3); ax3.legend(); ax3.axhline(0, linewidth=0.8)
            ax3.set_ylim(*pad_limits(y))
            st.pyplot(fig, clear_figure=True)
    except Exception as e:
        st.error(str(e))

if do_csv:
    try:
        # compute in SI
        L_si = L * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        if "UDL" in load_type:
            load_si = load_val * (14.5939 if units.startswith("Imperial") else (1e-3 if units.startswith("CGS") else 1.0))
        elif "Moment" in load_type:
            load_si = load_val * (1.355818 if units.startswith("Imperial") else (1e-7 if units.startswith("CGS") else 1.0))
        else:
            load_si = load_val * (4.4482216153 if units.startswith("Imperial") else (1e-5 if units.startswith("CGS") else 1.0))
        a_si = None if a is None else a * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L_si, load_si, E, I, c, mat_name, Sy, a_si)
        df = pd.DataFrame({"x_m": x, "V_N": V, "M_Nm": M, "delta_m": y})
        csv_bytes = df.to_csv(index=False).encode("utf-8")
        st.download_button("Download CSV (click to save)", csv_bytes, file_name="beam_results.csv", mime="text/csv")
    except Exception as e:
        st.error(str(e))

if do_pdf:
    try:
        # compute SI
        L_si = L * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        if "UDL" in load_type:
            load_si = load_val * (14.5939 if units.startswith("Imperial") else (1e-3 if units.startswith("CGS") else 1.0))
        elif "Moment" in load_type:
            load_si = load_val * (1.355818 if units.startswith("Imperial") else (1e-7 if units.startswith("CGS") else 1.0))
        else:
            load_si = load_val * (4.4482216153 if units.startswith("Imperial") else (1e-5 if units.startswith("CGS") else 1.0))
        a_si = None if a is None else a * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L_si, load_si, E, I, c, mat_name, Sy, a_si)
        pdf_buf = io.BytesIO()
        with PdfPages(pdf_buf) as pdf:
            # Cover / summary
            fig_sum = plt.figure(figsize=(8.5, 11))
            fig_sum.text(0.1,0.94,"Beam Analysis Report", fontsize=18)
            # Optional project fields
            proj = st.session_state.get("project", {})
            fig_sum.text(0.1,0.90,f"Project: {proj.get('name','—')}")
            fig_sum.text(0.55,0.90,f"Engineer: {proj.get('eng','—')}")
            fig_sum.text(0.1,0.87,f"Notes: {proj.get('notes','—')}")
            fig_sum.text(0.1,0.83,f"Beam: {beam.beam_type} | Load: {beam.load_type} | L = {beam.length:.3f} m")
            fig_sum.text(0.1,0.80,f"Material: {beam.material} | E = {((sci_notation(beam.E*10)+' dyn/cm²') if units.startswith('CGS') else (sci_notation(beam.E)+' Pa'))} | Sy = {((sci_notation(beam.yield_strength*10)+' dyn/cm²') if (beam.yield_strength and units.startswith('CGS')) else ((sci_notation(beam.yield_strength) if beam.yield_strength else '—') + (' Pa' if beam.yield_strength else '')))}")
            fig_sum.text(0.1,0.77,f"I = {sci_notation(beam.I)} m^4 | c = {sci_notation(beam.c)} m")
            Mabs = np.max(np.abs(M)); Vabs = np.max(np.abs(V)); yabs = np.max(np.abs(y))
            fig_sum.text(0.1,0.73,f"Max |V| = {sci_notation(Vabs)} N | Max |M| = {sci_notation(Mabs)} Nm | Max |δ| = {sci_notation(yabs)} m")
            if react:
                ry = 0.69; fig_sum.text(0.1, ry, "Reactions:", fontsize=12); ry -= 0.02
                for k, v in react.items(): fig_sum.text(0.12, ry, f"{k}: {sci_notation(v)}"); ry -= 0.018
            pdf.savefig(fig_sum); plt.close(fig_sum)

            # Plots page
            fig2 = plt.figure(figsize=(10, 10))
            fig2.suptitle(f"{beam.beam_type} – {beam.load_type}", fontsize=14)
            ax1 = fig2.add_subplot(3,1,1); ax1.plot(x, V, linewidth=2.2, label="V(x) [N]"); ax1.set_title("Shear V"); ax1.set_ylabel("N"); ax1.grid(True, alpha=0.3); ax1.legend(); ax1.axhline(0, linewidth=0.8); ax1.set_ylim(*pad_limits(V))
            ax2 = fig2.add_subplot(3,1,2); ax2.plot(x, M, linewidth=2.2, label="M(x) [Nm]"); ax2.set_title("Moment M"); ax2.set_ylabel("Nm"); ax2.grid(True, alpha=0.3); ax2.legend(); ax2.axhline(0, linewidth=0.8); ax2.set_ylim(*pad_limits(M))
            ax3 = fig2.add_subplot(3,1,3); ax3.plot(x, y, linewidth=2.2, label="δ(x) [m]"); ax3.set_title("Deflection δ"); ax3.set_xlabel("x (m)"); ax3.set_ylabel("m"); ax3.grid(True, alpha=0.3); ax3.legend(); ax3.axhline(0, linewidth=0.8); ax3.set_ylim(*pad_limits(y))
            plt.tight_layout(rect=[0, 0, 1, 0.96])
            pdf.savefig(fig2); plt.close(fig2)

        st.download_button("Download PDF (click to save)", pdf_buf.getvalue(), file_name="beam_report.pdf", mime="application/pdf")
    except Exception as e:
        st.error(str(e))

# ------------------------- Project details (for report) -------------------------
with st.expander("Project / Report details"):
    pn = st.text_input("Project name", key="proj_name")
    pe = st.text_input("Engineer", key="proj_eng")
    pt = st.text_area("Notes", key="proj_notes")
    st.session_state["project"] = {"name": pn, "eng": pe, "notes": pt}

# ------------------------- Export: Excel / Word -------------------------
if do_xlsx:
    try:
        # recompute SI
        L_si = L * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        if "UDL" in load_type:
            load_si = load_val * (14.5939 if units.startswith("Imperial") else (1e-3 if units.startswith("CGS") else 1.0))
        elif "Moment" in load_type:
            load_si = load_val * (1.355818 if units.startswith("Imperial") else (1e-7 if units.startswith("CGS") else 1.0))
        else:
            load_si = load_val * (4.4482216153 if units.startswith("Imperial") else (1e-5 if units.startswith("CGS") else 1.0))
        a_si = None if a is None else a * (0.3048 if units.startswith("Imperial") else (0.01 if units.startswith("CGS") else 1.0))
        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L_si, load_si, E, I, c, mat_name, Sy, a_si)
        df = pd.DataFrame({"x_m": x, "V_N": V, "M_Nm": M, "delta_m": y})
        import io as _io
        from pandas import ExcelWriter
        bio = _io.BytesIO()
        with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Diagrams")
            info = pd.DataFrame({
                "Parameter": ["Beam","Load","L (m)","Load","Material","E (Pa)","Sy (Pa)","I (m^4)","c (m)","Max|V| (N)","Max|M| (Nm)","Max|δ| (m)","Max σ (Pa)","n"],
                "Value": [
                    beam.beam_type, beam.load_type, f"{beam.length:.3f}", sci_notation(beam.load), beam.material,
                    sci_notation(E), (sci_notation(Sy) if Sy else "—"), sci_notation(I), sci_notation(c),
                    sci_notation(np.max(np.abs(V))), sci_notation(np.max(np.abs(M))), sci_notation(np.max(np.abs(y))),
                    (sci_notation(stress) if stress is not None else "N/A"), (f"{sf:.2f}" if sf else "—")
                ]
            })
            info.to_excel(writer, index=False, sheet_name="Summary")
        st.download_button("Download Excel (xlsx)", bio.getvalue(), file_name="beam_results.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        st.error(str(e))

if do_docx:
    try:
        # minimal docx export
        from docx import Document
        beam, x, V, M, y, stress, sf, react = compute_all(beam_type, load_type, L * (0.3048 if units.startswith("Imperial") else 1.0),
                                                          (load_val * (14.5939 if ("UDL" in load_type and units.startswith("Imperial")) else (1.355818 if ("Moment" in load_type and units.startswith("Imperial")) else (4.4482216153 if units.startswith("Imperial") else 1.0)))),
                                                          E, I, c, mat_name, Sy, (None if a is None else a * (0.3048 if units.startswith("Imperial") else 1.0)))
        doc = Document()
        p = st.session_state.get("project", {})
        doc.add_heading('Beam Analysis Report', 0)
        doc.add_paragraph(f"Project: {p.get('name','—')}  |  Engineer: {p.get('eng','—')}")
        doc.add_paragraph(f"Beam: {beam.beam_type} | Load: {beam.load_type} | L = {beam.length:.3f} m")
        doc.add_paragraph(f"Material: {beam.material} | E = {((sci_notation(E*10)+' dyn/cm²') if units.startswith('CGS') else (sci_notation(E)+' Pa'))} | Sy = {((sci_notation(Sy*10)+' dyn/cm²') if (Sy and units.startswith('CGS')) else ((sci_notation(Sy) if Sy else '—') + (' Pa' if Sy else '')))}")
        doc.add_paragraph(f"I = {sci_notation(I)} m^4 | c = {sci_notation(c)} m")
        doc.add_paragraph(f"Max |V| = {sci_notation(np.max(np.abs(V)))} N | Max |M| = {sci_notation(np.max(np.abs(M)))} Nm | Max |δ| = {sci_notation(np.max(np.abs(y)))} m")
        if stress is not None:
            doc.add_paragraph(f"Max σ = {sci_notation(stress)} Pa;  n = {(f'{sf:.2f}' if sf else '—')}")
        b = io.BytesIO(); doc.save(b)
        st.download_button("Download Word (.docx)", b.getvalue(), file_name="beam_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(str(e))


# ------------------------- Save / Load analysis as JSON -------------------------
import json
if 'project' not in st.session_state: st.session_state['project'] = {}
if do_save_json:
    payload = {
        "material": mat_name, "E": E, "Sy": Sy,
        "units": units, "beam_type": beam_type, "load_type": load_type,
        "L": L, "load_val": load_val, "a": a,
        "I": I, "c": c,
        "project": st.session_state.get("project", {}),
    }
    st.download_button("Download inputs.json", data=json.dumps(payload, indent=2), file_name="beam_inputs.json", mime="application/json")

if up_json is not None:
    try:
        data = json.load(up_json)
        st.success("Loaded inputs from JSON. Apply to fields if needed.")
        st.session_state['loaded_json'] = data
    except Exception as ex:
        st.error(f"Invalid JSON: {ex}")
# ------------------------- Help / parameter meanings -------------------------
with st.expander("Parameter meanings (full) & units"):
    st.markdown(
        """
        - **E (Young's Modulus)** — material stiffness (Pa)
        - **Sy (Yield Strength)** — stress for permanent deformation (Pa)
        - **L (Beam Length)** — total span (m) or ft (imperial mode)
        - **Load value** — `P` for point loads (N/lbf), `w` for UDL (N/m or lbf/ft), or applied **Moment** (Nm or lbf·ft)
        - **a (Load Position)** — distance from left support (m or ft)
        - **I (Area Moment of Inertia)** — bending resistance (m⁴)
        - **c (Extreme Fiber Distance)** — NA to outer fiber (m)
        - **V (Shear Force)** — internal shear (N or lbf)
        - **M (Bending Moment)** — internal bending (Nm or lbf·ft)
        - **δ (Deflection)** — vertical displacement (m or ft)
        - **σ (Bending Stress)** — outer fiber stress (Pa)
        - **n (Safety Factor)** — `Sy / σ` (dimensionless)
        """
    )

st.markdown("---")
st.caption("Alphabetical, searchable materials • Light/Dark toggle • Separate actions • Interactive plots optional • Exports: CSV, PDF, Excel, Word")
